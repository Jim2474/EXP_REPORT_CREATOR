from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


EXP_DIR = Path(__file__).resolve().parent
FIG_DIR = EXP_DIR / "figures"
IMG_DIR = EXP_DIR / "output_images"
LOG_PATH = EXP_DIR / "experiment3_output.txt"
REPORT_DIR = EXP_DIR / "reports"
REPORT_PATH = REPORT_DIR / "2300810617李俊明_实验三_A4示例报告.docx"


EXERCISE_TITLES = {
    "3-1": "练习 3-1：Bode图绘制",
    "3-2": "练习 3-2：稳定裕度计算",
    "3-3": "练习 3-3：Nyquist图绘制",
    "3-4": "练习 3-4：根轨迹绘制",
    "3-5": "练习 3-5：不同增益下的阶跃响应",
    "3-6": "练习 3-6：频率响应数据表",
    "3-7": "练习 3-7：Nyquist图与稳定性判断",
    "3-8": "练习 3-8：多系统根轨迹对比",
}


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\consolab.ttf" if bold else r"C:\Windows\Fonts\consola.ttf"),
        Path(r"C:\Windows\Fonts\courbd.ttf" if bold else r"C:\Windows\Fonts\cour.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def split_log_sections(text: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current_key: str | None = None
    for line in text.splitlines():
        m = re.match(r"^===== Exercise 3-(\d+) =====$", line.strip())
        if m:
            current_key = f"3-{m.group(1)}"
            sections[current_key] = [f">> % {EXERCISE_TITLES[current_key]}"]
            continue
        if current_key:
            sections[current_key].append(line.rstrip())
    return {key: "\n".join(lines).strip() for key, lines in sections.items()}


def render_matlab_style(text: str, output_path: Path, title: str) -> None:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    font = load_font(25)
    font_bold = load_font(25, bold=True)
    title_font = load_font(24, bold=True)

    lines = text.splitlines()
    max_chars = 92
    wrapped: list[str] = []
    for line in lines:
        if len(line) <= max_chars:
            wrapped.append(line)
            continue
        indent = "    " if line.startswith(" ") else ""
        rest = line
        while len(rest) > max_chars:
            wrapped.append(rest[:max_chars])
            rest = indent + rest[max_chars:]
        wrapped.append(rest)

    margin_x = 34
    margin_y = 26
    line_h = 34
    width = 1500
    height = margin_y * 2 + 40 + max(1, len(wrapped)) * line_h
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    draw.rectangle((0, 0, width - 1, height - 1), outline=(205, 205, 205), width=2)
    draw.rectangle((0, 0, width - 1, 38), fill=(245, 247, 250), outline=(215, 215, 215))
    draw.text((14, 7), title, font=title_font, fill=(55, 55, 55))

    y = margin_y + 34
    for line in wrapped:
        color = (0, 0, 0)
        use_font = font
        if line.startswith(">>"):
            color = (0, 0, 170)
            use_font = font_bold
        elif line.startswith("图形已保存"):
            color = (0, 120, 0)
        draw.text((margin_x, y), line, font=use_font, fill=color)
        y += line_h

    image.save(output_path)


def set_font(run, name: str = "宋体", size: float = 11, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")


def add_heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def add_paragraphs(doc: Document, paragraphs: Iterable[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(text)


def add_result_table(doc: Document, entries: list[tuple[str, Path]]) -> None:
    for start in range(0, len(entries), 4):
        chunk = entries[start : start + 4]
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                cell.width = Cm(8.05)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
        for pos, (caption_text, image_path) in enumerate(chunk):
            cell = table.cell(pos // 2, pos % 2)
            caption = cell.paragraphs[0]
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.first_line_indent = Cm(0)
            caption.paragraph_format.space_after = Pt(2)
            run = caption.add_run(caption_text)
            set_font(run, "黑体", 8.5, True)

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            shape = p.add_run().add_picture(str(image_path))
            max_w = Cm(7.45)
            max_h = Cm(8.0)
            scale = min(max_w / shape.width, max_h / shape.height)
            shape.width = int(shape.width * scale)
            shape.height = int(shape.height * scale)
        if start + 4 < len(entries):
            doc.add_page_break()


def build_report(entries: list[tuple[str, Path]]) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.35
    styles["Normal"].paragraph_format.first_line_indent = Cm(0.74)
    styles["Normal"].paragraph_format.space_after = Pt(0)
    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("实验三  线性控制系统的频域响应分析")
    set_font(run, "黑体", 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("姓名：李俊明    学号：2300810617")
    set_font(run, "宋体", 11)

    add_heading(doc, "一、实验目的")
    add_paragraphs(
        doc,
        [
            "本次实验主要熟悉MATLAB系统中有关频域分析的命令，包括bode、margin、nyquist和rlocus等函数的使用方法；掌握线性系统频率响应分析的一般方法，包括Bode图绘制、稳定裕度计算、Nyquist曲线绘制和根轨迹分析。通过实验加深对控制系统频域特性的理解。",
        ],
    )

    add_heading(doc, "二、实验环境")
    add_paragraphs(
        doc,
        [
            "实验软件：MATLAB R2024b。",
            "实验工具箱：Control System Toolbox。",
            "实验内容：Bode图绘制、稳定裕度计算、Nyquist图绘制、根轨迹分析、频率响应与稳定性判断。",
        ],
    )

    add_heading(doc, "三、实验内容")
    add_paragraphs(
        doc,
        [
            "1. 使用bode命令绘制开环系统的Bode图，分析系统的频率响应特性。",
            "2. 使用margin命令计算系统的幅值裕度和相角裕度，判断系统稳定性。",
            "3. 使用nyquist命令绘制系统的Nyquist图，分析系统的稳定性。",
            "4. 使用rlocus命令绘制系统的根轨迹，分析增益变化对系统稳定性的影响。",
            "5. 通过改变开环增益，观察闭环系统阶跃响应的变化规律。",
            "6. 综合运用频域分析方法，判断闭环系统的稳定性。",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "四、实验结果")
    p = doc.add_paragraph("以下结果由MATLAB批处理实际运行得到。命令行类结果采用MATLAB命令窗口风格渲染，曲线图由MATLAB直接保存。")
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_result_table(doc, entries)

    doc.add_page_break()
    add_heading(doc, "五、实验结果与分析")
    add_paragraphs(
        doc,
        [
            "从实验结果可以看出，bode命令能够方便地绘制系统的Bode图，显示幅频特性和相频特性。margin命令可以快速计算系统的幅值裕度和相角裕度，为判断系统稳定性提供定量依据。",
            "Nyquist图通过绘制系统频率特性的极坐标图，可以直观地判断闭环系统的稳定性。根据Nyquist稳定性判据，当Nyquist曲线不包围(-1,j0)点时，闭环系统稳定。",
            "根轨迹分析表明，随着开环增益的增大，闭环极点会从开环极点出发沿根轨迹移动。当增益超过临界值时，极点进入右半平面，系统变得不稳定。实验中不同增益下的阶跃响应验证了这一结论：增益越大，超调量越大，系统稳定性越差。",
        ],
    )

    add_heading(doc, "六、实验总结")
    add_paragraphs(
        doc,
        [
            "通过本次实验，我熟悉了MATLAB在控制系统频域分析中的基本用法，掌握了bode、margin、nyquist和rlocus等常用命令。实验说明MATLAB能够有效简化频域分析过程，并能通过图形方式直观展示系统的频率特性。本次实验为后续控制系统校正设计打下了基础。",
        ],
    )

    doc.save(REPORT_PATH)


def main() -> None:
    log_text = LOG_PATH.read_text(encoding="utf-8", errors="replace")
    sections = split_log_sections(log_text)

    entries: list[tuple[str, Path]] = []

    # Exercise 3-1: Bode图
    image_path = IMG_DIR / "exercise_3_1_text.png"
    render_matlab_style(sections["3-1"], image_path, EXERCISE_TITLES["3-1"])
    entries.append((EXERCISE_TITLES["3-1"], image_path))
    entries.append(("练习 3-1：Bode图", FIG_DIR / "exercise_3_1.png"))

    # Exercise 3-2: 稳定裕度
    image_path = IMG_DIR / "exercise_3_2_text.png"
    render_matlab_style(sections["3-2"], image_path, EXERCISE_TITLES["3-2"])
    entries.append((EXERCISE_TITLES["3-2"], image_path))
    entries.append(("练习 3-2：稳定裕度Bode图", FIG_DIR / "exercise_3_2.png"))

    # Exercise 3-3: Nyquist图
    image_path = IMG_DIR / "exercise_3_3_text.png"
    render_matlab_style(sections["3-3"], image_path, EXERCISE_TITLES["3-3"])
    entries.append((EXERCISE_TITLES["3-3"], image_path))
    entries.append(("练习 3-3：Nyquist图", FIG_DIR / "exercise_3_3.png"))

    # Exercise 3-4: 根轨迹
    image_path = IMG_DIR / "exercise_3_4_text.png"
    render_matlab_style(sections["3-4"], image_path, EXERCISE_TITLES["3-4"])
    entries.append((EXERCISE_TITLES["3-4"], image_path))
    entries.append(("练习 3-4：根轨迹", FIG_DIR / "exercise_3_4.png"))

    # Exercise 3-5: 阶跃响应
    image_path = IMG_DIR / "exercise_3_5_text.png"
    render_matlab_style(sections["3-5"], image_path, EXERCISE_TITLES["3-5"])
    entries.append((EXERCISE_TITLES["3-5"], image_path))
    entries.append(("练习 3-5：不同增益下的阶跃响应", FIG_DIR / "exercise_3_5.png"))

    # Exercise 3-6: 频率响应数据表
    image_path = IMG_DIR / "exercise_3_6_text.png"
    render_matlab_style(sections["3-6"], image_path, EXERCISE_TITLES["3-6"])
    entries.append((EXERCISE_TITLES["3-6"], image_path))
    entries.append(("练习 3-6：稳定裕度Bode图", FIG_DIR / "exercise_3_6.png"))

    # Exercise 3-7: Nyquist图与稳定性判断
    image_path = IMG_DIR / "exercise_3_7_text.png"
    render_matlab_style(sections["3-7"], image_path, EXERCISE_TITLES["3-7"])
    entries.append((EXERCISE_TITLES["3-7"], image_path))
    entries.append(("练习 3-7：Nyquist图", FIG_DIR / "exercise_3_7.png"))

    # Exercise 3-8: 多系统根轨迹对比
    image_path = IMG_DIR / "exercise_3_8_text.png"
    render_matlab_style(sections["3-8"], image_path, EXERCISE_TITLES["3-8"])
    entries.append((EXERCISE_TITLES["3-8"], image_path))
    entries.append(("练习 3-8：根轨迹对比", FIG_DIR / "exercise_3_8.png"))

    build_report(entries)
    print(f"report={REPORT_PATH}")
    print(f"entries={len(entries)}")


if __name__ == "__main__":
    main()
