from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


EXP_DIR = Path(__file__).resolve().parent
FIG_DIR = EXP_DIR / "figures"
IMG_DIR = EXP_DIR / "output_images"
LOG_PATH = EXP_DIR / "experiment2_output.txt"
REPORT_DIR = EXP_DIR / "reports"
REPORT_PATH = REPORT_DIR / "2300810617李俊明_实验二_A4示例报告_修正版.docx"


EXERCISE_TITLES = {
    "2-1": "练习 2-1：求多项式的根",
    "2-2": "练习 2-2：由根求多项式",
    "2-3": "练习 2-3：部分分式展开",
    "2-4": "练习 2-4：零极点增益模型",
    "2-5": "练习 2-5：闭环传递函数",
    "2-6": "练习 2-6：电压、电流和功率曲线",
    "2-7": "练习 2-7：函数曲线及极值",
    "2-8": "练习 2-8：阶跃响应指标函数调用",
    "2-9": "练习 2-9：典型系统阶跃响应指标",
    "2-10": "练习 2-10：单位阶跃响应曲线及指标",
}


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\consolab.ttf" if bold else r"C:\Windows\Fonts\consola.ttf"),
        Path(r"C:\Windows\Fonts\courbd.ttf" if bold else r"C:\Windows\Fonts\cour.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def split_log_sections(text: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current_key: str | None = None
    for line in text.splitlines():
        m = re.match(r"^===== Exercise 2-(\d+) =====$", line.strip())
        if m:
            current_key = f"2-{m.group(1)}"
            sections[current_key] = [f">> % {EXERCISE_TITLES[current_key]}"]
            continue
        if current_key:
            sections[current_key].append(line.rstrip())
    return {key: "\n".join(lines).strip() for key, lines in sections.items()}


def render_matlab_style(text: str, output_path: Path, title: str) -> None:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    font = load_font(25)
    font_bold = load_font(25, bold=True)
    title_font = load_font(24, bold=True)

    lines = text.splitlines()
    # Keep each result compact enough for A4 while retaining readable text.
    max_chars = 92
    wrapped: list[str] = []
    for line in lines:
        if len(line) <= max_chars:
            wrapped.append(line)
            continue
        indent = "    " if line.startswith(" ") else ""
        rest = line
        while len(rest) > max_chars:
            wrapped.append(rest[:max_chars])
            rest = indent + rest[max_chars:]
        wrapped.append(rest)

    margin_x = 34
    margin_y = 26
    line_h = 34
    width = 1500
    height = margin_y * 2 + 40 + max(1, len(wrapped)) * line_h
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    # Thin MATLAB-like command-window frame.
    draw.rectangle((0, 0, width - 1, height - 1), outline=(205, 205, 205), width=2)
    draw.rectangle((0, 0, width - 1, 38), fill=(245, 247, 250), outline=(215, 215, 215))
    draw.text((14, 7), title, font=title_font, fill=(55, 55, 55))

    y = margin_y + 34
    for line in wrapped:
        color = (0, 0, 0)
        use_font = font
        if line.startswith(">>"):
            color = (0, 0, 170)
            use_font = font_bold
        elif line.startswith("图形已保存"):
            color = (0, 120, 0)
        draw.text((margin_x, y), line, font=use_font, fill=color)
        y += line_h

    image.save(output_path)


def set_font(run, name: str = "宋体", size: float = 11, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")


def add_heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def add_paragraphs(doc: Document, paragraphs: Iterable[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(text)


def add_result_table(doc: Document, entries: list[tuple[str, Path]]) -> None:
    for start in range(0, len(entries), 4):
        chunk = entries[start : start + 4]
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                cell.width = Cm(8.05)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
        for pos, (caption_text, image_path) in enumerate(chunk):
            cell = table.cell(pos // 2, pos % 2)
            caption = cell.paragraphs[0]
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.first_line_indent = Cm(0)
            caption.paragraph_format.space_after = Pt(2)
            run = caption.add_run(caption_text)
            set_font(run, "黑体", 8.5, True)

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            shape = p.add_run().add_picture(str(image_path))
            max_w = Cm(7.45)
            max_h = Cm(8.0)
            scale = min(max_w / shape.width, max_h / shape.height)
            shape.width = int(shape.width * scale)
            shape.height = int(shape.height * scale)
        if start + 4 < len(entries):
            doc.add_page_break()


def build_report(entries: list[tuple[str, Path]]) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.35
    styles["Normal"].paragraph_format.first_line_indent = Cm(0.74)
    styles["Normal"].paragraph_format.space_after = Pt(0)
    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("实验二  控制系统的数学模型及时域分析")
    set_font(run, "黑体", 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("姓名：李俊明    学号：2300810617")
    set_font(run, "宋体", 11)

    add_heading(doc, "一、实验目的")
    add_paragraphs(
        doc,
        [
            "本次实验主要熟悉 MATLAB 线性控制系统模型的基本描述方法，掌握传递函数模型、零极点增益模型和状态空间模型之间的转换方法；学习使用 roots、poly、residue、tf2zp、feedback、step 等命令完成控制系统建模与时域分析。通过绘制函数曲线和阶跃响应曲线，加深对控制系统动态性能指标的理解。",
        ],
    )

    add_heading(doc, "二、实验环境")
    add_paragraphs(
        doc,
        [
            "实验软件：MATLAB R2024b。",
            "实验工具箱：Control System Toolbox。",
            "实验内容：多项式运算、部分分式展开、数学模型转换、闭环传递函数求取、函数作图及阶跃响应分析。",
        ],
    )

    add_heading(doc, "三、实验内容")
    add_paragraphs(
        doc,
        [
            "1. 使用 roots 命令求多项式的根，并使用 poly 命令由给定根求多项式系数。",
            "2. 使用 residue 命令对给定有理分式进行部分分式展开。",
            "3. 使用 tf2zp 命令求传递函数的零点、极点和增益。",
            "4. 使用 feedback 命令求由两个子系统构成的负反馈闭环系统传递函数。",
            "5. 使用 plot、subplot 等命令绘制电压、电流、功率以及给定函数曲线。",
            "6. 编写并调用阶跃响应特征参数函数，计算超调量、上升时间、峰值时间和调整时间。",
            "7. 使用 step 命令绘制单位阶跃响应曲线，并记录系统动态性能指标。",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "四、实验结果")
    p = doc.add_paragraph("以下结果由 MATLAB 批处理实际运行得到。命令行类结果采用 MATLAB 命令窗口风格渲染，曲线图由 MATLAB 直接保存。")
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_result_table(doc, entries)

    doc.add_page_break()
    add_heading(doc, "五、实验结果与分析")
    add_paragraphs(
        doc,
        [
            "从实验结果可以看出，roots 与 poly 能够方便地实现多项式根和多项式系数之间的转换；residue 命令可以将有理分式展开为部分分式形式，便于后续进行系统响应分析。",
            "在数学模型转换中，tf2zp 可以直接得到系统的零点、极点和增益。反馈连接实验表明，MATLAB 的 feedback 命令能够根据子系统传递函数快速求出闭环系统模型，避免了手工推导高阶多项式时容易出现的计算错误。",
            "在时域分析部分，step 命令可以得到系统单位阶跃响应曲线。通过编程计算超调量、上升时间、峰值时间和调整时间，可以更直观地比较不同系统的动态性能。实验中不同传递函数的阶跃响应差异明显，说明系统参数会直接影响响应速度、超调和稳定过程。",
        ],
    )

    add_heading(doc, "六、实验总结")
    add_paragraphs(
        doc,
        [
            "通过本次实验，我熟悉了 MATLAB 在控制系统数学模型建立和时域分析中的基本用法，掌握了多项式求根、部分分式展开、模型转换、反馈连接和阶跃响应分析等常用命令。实验说明 MATLAB 能够有效简化控制系统计算过程，并能通过图形方式直观展示系统动态特性。本次实验为后续频域分析、Simulink 建模以及控制系统校正设计打下了基础。",
        ],
    )

    doc.save(REPORT_PATH)


def main() -> None:
    log_text = LOG_PATH.read_text(encoding="utf-8", errors="replace")
    sections = split_log_sections(log_text)

    entries: list[tuple[str, Path]] = []
    for key in ["2-1", "2-2", "2-3", "2-4", "2-5"]:
        image_path = IMG_DIR / f"exercise_{key.replace('-', '_')}.png"
        render_matlab_style(sections[key], image_path, EXERCISE_TITLES[key])
        entries.append((EXERCISE_TITLES[key], image_path))

    entries.append((EXERCISE_TITLES["2-6"], FIG_DIR / "exercise_2_6.png"))

    image_path = IMG_DIR / "exercise_2_7_text.png"
    render_matlab_style(sections["2-7"], image_path, EXERCISE_TITLES["2-7"])
    entries.append(("练习 2-7：极值计算结果", image_path))
    entries.append(("练习 2-7：函数曲线", FIG_DIR / "exercise_2_7.png"))

    for key in ["2-8", "2-9"]:
        image_path = IMG_DIR / f"exercise_{key.replace('-', '_')}.png"
        render_matlab_style(sections[key], image_path, EXERCISE_TITLES[key])
        entries.append((EXERCISE_TITLES[key], image_path))

    entries.append(("练习 2-9：阶跃响应曲线", FIG_DIR / "exercise_2_9.png"))

    image_path = IMG_DIR / "exercise_2_10_text.png"
    render_matlab_style(sections["2-10"], image_path, EXERCISE_TITLES["2-10"])
    entries.append(("练习 2-10：动态指标计算结果", image_path))
    entries.append(("练习 2-10：单位阶跃响应曲线", FIG_DIR / "exercise_2_10.png"))

    build_report(entries)
    print(f"report={REPORT_PATH}")
    print(f"entries={len(entries)}")


if __name__ == "__main__":
    main()
