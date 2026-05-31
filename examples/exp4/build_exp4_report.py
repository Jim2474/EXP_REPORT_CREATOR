from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


EXP_DIR = Path(__file__).resolve().parent
FIG_DIR = EXP_DIR / "figures"
IMG_DIR = EXP_DIR / "output_images"
LOG_PATH = EXP_DIR / "experiment4_output.txt"
REPORT_DIR = EXP_DIR / "reports"
REPORT_PATH = REPORT_DIR / "2300810617李俊明_实验四_A4示例报告.docx"


EXERCISE_TITLES = {
    "4-1": "练习 4-1：T1、T2、T3 阶跃响应",
    "4-2-1": "练习 4-2-①：sigma变化阶跃响应",
    "4-2-2": "练习 4-2-②：theta变化阶跃响应",
    "4-3-1": "练习 4-3-①：n(s)=1.5 阶跃响应",
    "4-3-2": "练习 4-3-②：右平面零点响应",
    "4-3-3": "练习 4-3-③：左平面零点响应",
    "4-3-4": "练习 4-3-④：结果汇总",
}

# 使用 SIMULINK 生成的图片
USE_SIMULINK_FIGURES = True


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\consolab.ttf" if bold else r"C:\Windows\Fonts\consola.ttf"),
        Path(r"C:\Windows\Fonts\courbd.ttf" if bold else r"C:\Windows\Fonts\cour.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def split_log_sections(text: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current_key: str | None = None
    for line in text.splitlines():
        m = re.match(r"^===== Exercise 4-(\d+) =====$", line.strip())
        if m:
            current_key = f"4-{m.group(1)}"
            sections[current_key] = [f">> % {EXERCISE_TITLES.get(current_key, f'练习 {current_key}')}"]
            continue
        if current_key:
            sections[current_key].append(line.rstrip())
    return {key: "\n".join(lines).strip() for key, lines in sections.items()}


def render_matlab_style(text: str, output_path: Path, title: str) -> None:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    font = load_font(25)
    font_bold = load_font(25, bold=True)
    title_font = load_font(24, bold=True)

    lines = text.splitlines()
    max_chars = 92
    wrapped: list[str] = []
    for line in lines:
        if len(line) <= max_chars:
            wrapped.append(line)
            continue
        indent = "    " if line.startswith(" ") else ""
        rest = line
        while len(rest) > max_chars:
            wrapped.append(rest[:max_chars])
            rest = indent + rest[max_chars:]
        wrapped.append(rest)

    margin_x = 34
    margin_y = 26
    line_h = 34
    width = 1500
    height = margin_y * 2 + 40 + max(1, len(wrapped)) * line_h
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    draw.rectangle((0, 0, width - 1, height - 1), outline=(205, 205, 205), width=2)
    draw.rectangle((0, 0, width - 1, 38), fill=(245, 247, 250), outline=(215, 215, 215))
    draw.text((14, 7), title, font=title_font, fill=(55, 55, 55))

    y = margin_y + 34
    for line in wrapped:
        color = (0, 0, 0)
        use_font = font
        if line.startswith(">>"):
            color = (0, 0, 170)
            use_font = font_bold
        elif line.startswith("图形已保存"):
            color = (0, 120, 0)
        draw.text((margin_x, y), line, font=use_font, fill=color)
        y += line_h

    image.save(output_path)


def set_font(run, name: str = "宋体", size: float = 11, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")


def add_heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def add_paragraphs(doc: Document, paragraphs: Iterable[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(text)


def add_result_table(doc: Document, entries: list[tuple[str, Path]]) -> None:
    for start in range(0, len(entries), 4):
        chunk = entries[start : start + 4]
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                cell.width = Cm(8.05)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
        for pos, (caption_text, image_path) in enumerate(chunk):
            cell = table.cell(pos // 2, pos % 2)
            caption = cell.paragraphs[0]
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.first_line_indent = Cm(0)
            caption.paragraph_format.space_after = Pt(2)
            run = caption.add_run(caption_text)
            set_font(run, "黑体", 8.5, True)

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            shape = p.add_run().add_picture(str(image_path))
            max_w = Cm(7.45)
            max_h = Cm(8.0)
            scale = min(max_w / shape.width, max_h / shape.height)
            shape.width = int(shape.width * scale)
            shape.height = int(shape.height * scale)
        if start + 4 < len(entries):
            doc.add_page_break()


def build_report(entries: list[tuple[str, Path]], log_sections: dict[str, str]) -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.35
    styles["Normal"].paragraph_format.first_line_indent = Cm(0.74)
    styles["Normal"].paragraph_format.space_after = Pt(0)
    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("实验四  SIMULINK 基本用法")
    set_font(run, "黑体", 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("姓名：李俊明    学号：2300810617")
    set_font(run, "宋体", 11)

    add_heading(doc, "一、实验目的")
    add_paragraphs(
        doc,
        [
            "本次实验学习 SIMULINK 软件工具的使用方法，掌握 SIMULINK 模块库的基本功能和操作方法；用 SIMULINK 仿真线性系统，观察不同参数下系统的阶跃响应特性；学习用 SIMULINK 软件工具对经典控制系统进行仿真设计的基本方法，包括模块搭建、参数设置和仿真运行。",
        ],
    )

    add_heading(doc, "二、实验环境")
    add_paragraphs(
        doc,
        [
            "实验软件：MATLAB R2024b + SIMULINK。",
            "实验内容：SIMULINK 模块库使用、系统模型搭建、阶跃响应仿真、二阶欠阻尼系统分析、非最小相位系统分析。",
        ],
    )

    add_heading(doc, "三、实验内容")
    add_paragraphs(
        doc,
        [
            "1. 在 SIMULINK 环境下搭建 T1、T2、T3 系统模型，观察并比较三个系统的阶跃响应。",
            "2. 对典型二阶欠阻尼系统，分别改变 sigma 和 theta 参数，分析参数变化对阶跃响应的影响。",
            "3. 对非最小相位系统，研究右平面零点和左平面零点对阶跃响应的不同作用。",
            "4. 记录各系统的超调量、峰值时间和过渡过程时间等动态性能指标。",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "四、实验结果")

    # Exercise 4-1 results
    add_heading(doc, "练习 4-1：T1、T2、T3 系统的阶跃响应")
    add_paragraphs(
        doc,
        [
            "三个系统的传递函数分别为：",
            "T1(s) = 1/(s² + 2s + 1) = 1/(s+1)²",
            "T2(s) = 1/(s² + 2s + 2)",
            "T3(s) = 1/(s² + 2s + 3)",
            "三个系统均为二阶系统，但阻尼特性不同。T1 为临界阻尼（ζ=1），T2 和 T3 为欠阻尼系统。",
        ],
    )
    entries_41 = [(EXERCISE_TITLES["4-1"], FIG_DIR / "exercise_4_1.png")]
    add_result_table(doc, entries_41)

    doc.add_page_break()
    # Exercise 4-2 results
    add_heading(doc, "练习 4-2：典型二阶欠阻尼系统仿真")
    add_paragraphs(
        doc,
        [
            "典型二阶欠阻尼系统传递函数：G(s) = ωn²/(s² + 2ζωn·s + ωn²)",
            "极点位置：s = -σ ± jωa，其中 σ = ζωn，ωa = ωn√(1-ζ²)，cos(θ) = ζ",
        ],
    )

    add_heading(doc, "4-2-①：固定 ωa=1，改变 σ")
    add_paragraphs(
        doc,
        [
            "当 σ = 0.5 时，ωn = √1.25 ≈ 1.118，ζ ≈ 0.447，系统欠阻尼，有明显超调。",
            "当 σ = 1 时，ωn = √2 ≈ 1.414，ζ ≈ 0.707，系统仍欠阻尼，超调减小。",
            "当 σ = 5 时，ωn = √26 ≈ 5.099，ζ ≈ 0.981，系统接近临界阻尼，超调很小。",
            "随着 σ 增大，阻尼比 ζ 增大，超调量减小，响应速度加快但振荡减弱。",
        ],
    )
    entries_42_1 = [(EXERCISE_TITLES["4-2-1"], FIG_DIR / "exercise_4_2_1.png")]
    add_result_table(doc, entries_42_1)

    doc.add_page_break()
    add_heading(doc, "4-2-②：固定 ωn=2，改变 θ")
    add_paragraphs(
        doc,
        [
            "当 θ = 30° 时，ζ = cos(30°) ≈ 0.866，系统接近临界阻尼，超调很小。",
            "当 θ = 45° 时，ζ = cos(45°) ≈ 0.707，系统欠阻尼，有适度超调。",
            "当 θ = 60° 时，ζ = cos(60°) = 0.5，系统欠阻尼，超调较大。",
            "随着 θ 增大，阻尼比 ζ 减小，超调量增大，振荡加剧。",
        ],
    )
    entries_42_2 = [(EXERCISE_TITLES["4-2-2"], FIG_DIR / "exercise_4_2_2.png")]
    add_result_table(doc, entries_42_2)

    doc.add_page_break()
    # Exercise 4-3 results
    add_heading(doc, "练习 4-3：非最小相位系统仿真")
    add_paragraphs(
        doc,
        [
            "基础传递函数：G(s) = 1.5/(s² + s + 1)",
            "研究非最小相位零点（右平面零点）和最小相位零点（左平面零点）对系统响应的影响。",
        ],
    )

    add_heading(doc, "4-3-①：n(s) = 1.5 的阶跃响应")
    entries_43_1 = [(EXERCISE_TITLES["4-3-1"], FIG_DIR / "exercise_4_3_1.png")]
    add_result_table(doc, entries_43_1)

    doc.add_page_break()
    add_heading(doc, "4-3-②：右平面零点 n(s) = (-s+a)/a")
    add_paragraphs(
        doc,
        [
            "右平面零点（RHP zero）会使系统响应出现反向超调，即响应初期会向相反方向偏移。",
            "随着 a 值增大，零点远离原点，反向超调效应减弱。",
        ],
    )
    entries_43_2 = [(EXERCISE_TITLES["4-3-2"], FIG_DIR / "exercise_4_3_2.png")]
    add_result_table(doc, entries_43_2)

    doc.add_page_break()
    add_heading(doc, "4-3-③：左平面零点 n(s) = (s+a)/a")
    add_paragraphs(
        doc,
        [
            "左平面零点（LHP zero）会加速系统响应，增大超调量。",
            "随着 a 值增大，零点远离原点，对响应的影响减弱。",
        ],
    )
    entries_43_3 = [(EXERCISE_TITLES["4-3-3"], FIG_DIR / "exercise_4_3_3.png")]
    add_result_table(doc, entries_43_3)

    doc.add_page_break()
    add_heading(doc, "4-3-④：结果汇总与分析")
    add_paragraphs(
        doc,
        [
            "右平面零点（非最小相位零点）的作用：使系统响应出现初始反向偏移，增加调节时间，降低系统性能。零点越靠近原点（a 越小），反向效应越明显。",
            "左平面零点（最小相位零点）的作用：加速系统响应，增大超调量，但不会引起反向偏移。零点越靠近原点（a 越小），加速效应越明显。",
            "非最小相位系统的控制难度较大，因为其响应特性与常规系统相反，设计控制器时需要特别注意。",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "五、实验结果与分析")
    add_paragraphs(
        doc,
        [
            "从练习 4-1 的结果可以看出，T1、T2、T3 三个系统的阻尼特性不同，导致阶跃响应差异明显。T1 为临界阻尼系统，响应无超调；T2 和 T3 为欠阻尼系统，有不同程度的超调和振荡。",
            "练习 4-2 的结果表明，二阶系统的动态性能主要由阻尼比 ζ 和自然频率 ωn 决定。增大 σ（即增大 ζ）会减小超调量，增大 θ（即减小 ζ）会增大超调量。这些规律与理论分析一致。",
            "练习 4-3 的结果验证了非最小相位零点的特殊性质。右平面零点会使响应出现反向超调，这是非最小相位系统的典型特征。左平面零点则会加速响应，增大超调量。在实际控制系统设计中，应尽量避免非最小相位特性，或采取特殊控制策略来处理。",
        ],
    )

    add_heading(doc, "六、实验总结")
    add_paragraphs(
        doc,
        [
            "通过本次实验，我学习了 SIMULINK 的基本使用方法，掌握了模块库的分类和功能模块的操作方法。通过搭建不同系统的仿真模型，观察了参数变化对系统响应的影响，加深了对二阶系统动态特性的理解。",
            "实验还验证了非最小相位零点的特殊作用，认识到右平面零点会使系统响应出现反向偏移，这对控制系统设计具有重要指导意义。本次实验为后续控制系统校正设计打下了基础。",
        ],
    )

    doc.save(REPORT_PATH)


def main() -> None:
    log_text = LOG_PATH.read_text(encoding="utf-8", errors="replace")
    sections = split_log_sections(log_text)

    # Render text output images for sections with command-line output
    entries: list[tuple[str, Path]] = []

    # Determine figure suffix based on whether SIMULINK figures exist
    simulink_suffix = "_simulink"
    use_sim = USE_SIMULINK_FIGURES and (FIG_DIR / "exercise_4_1_simulink.png").exists()

    # Exercise 4-1 text
    if "4-1" in sections:
        image_path = IMG_DIR / "exercise_4_1_text.png"
        render_matlab_style(sections["4-1"], image_path, EXERCISE_TITLES["4-1"])
        entries.append((EXERCISE_TITLES["4-1"], image_path))
    # Exercise 4-1 figure
    fig_suffix = simulink_suffix if use_sim else ""
    entries.append((EXERCISE_TITLES["4-1"], FIG_DIR / f"exercise_4_1{fig_suffix}.png"))

    # Exercise 4-2 text sections
    for sub in ["1", "2"]:
        key = f"4-2-{sub}"
        if key in sections:
            image_path = IMG_DIR / f"exercise_4_2_{sub}_text.png"
            render_matlab_style(sections[key], image_path, EXERCISE_TITLES[key])
            entries.append((EXERCISE_TITLES[key], image_path))

    # Exercise 4-2 figures
    entries.append((EXERCISE_TITLES["4-2-1"], FIG_DIR / f"exercise_4_2_1{fig_suffix}.png"))
    entries.append((EXERCISE_TITLES["4-2-2"], FIG_DIR / f"exercise_4_2_2{fig_suffix}.png"))

    # Exercise 4-3 text sections
    for sub in ["1", "2", "3", "4"]:
        key = f"4-3-{sub}"
        if key in sections:
            image_path = IMG_DIR / f"exercise_4_3_{sub}_text.png"
            render_matlab_style(sections[key], image_path, EXERCISE_TITLES[key])
            entries.append((EXERCISE_TITLES[key], image_path))

    # Exercise 4-3 figures (4-3-1 doesn't have simulink version, use original)
    entries.append((EXERCISE_TITLES["4-3-1"], FIG_DIR / "exercise_4_3_1.png"))
    entries.append((EXERCISE_TITLES["4-3-2"], FIG_DIR / f"exercise_4_3_2{fig_suffix}.png"))
    entries.append((EXERCISE_TITLES["4-3-3"], FIG_DIR / f"exercise_4_3_3{fig_suffix}.png"))

    build_report(entries, sections)
    print(f"report={REPORT_PATH}")
    print(f"entries={len(entries)}")


if __name__ == "__main__":
    main()
